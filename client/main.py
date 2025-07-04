import base64
import streamlit as st
import grpc
import time
import LMS_pb2
import LMS_pb2_grpc
import jwt
import threading
import json
import io
from docx import Document



# servers =  [
#         "172.17.49.41:50054",
#         "172.17.49.41:50055",
#         "172.17.49.189:50051",
#         "172.17.49.189:50052",
#         "172.17.49.189:50053"
# ]

servers = [
     "127.0.0.1:50051",
        "127.0.0.1:50054",
        "127.0.0.1:50055",
        "127.0.0.1:50056",
        "127.0.0.1:50053"
]
placeholder = st.empty()
querieslist = st.empty()

# Initialize global state
if 'leader' not in st.session_state:
    st.session_state['leader'] = 'none'
if 'token' not in st.session_state:
    st.session_state['token'] = None
if 'role' not in st.session_state:
    st.session_state['role'] = None
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'channel' not in st.session_state:
    st.session_state['channel'] = None
if 'stub' not in st.session_state:
    st.session_state['stub'] = None



def display_leader_status():
    status_placeholder = st.empty()
    status_placeholder.markdown(
        f"<div style='text-align: right; font-weight: bold;'>Leader : {st.session_state['leader']}</div>",
        unsafe_allow_html=True
    )


def verify_token(token):
    try:
        decoded_payload = jwt.decode(token, "SECRET", algorithms=["HS256"])
        print(decoded_payload)
        st.session_state['role'] = decoded_payload['role']
        st.session_state['username'] = decoded_payload['username']
        return 
    except jwt.ExpiredSignatureError:
        print("JWT expired")
        return "Invalid"
    except jwt.InvalidTokenError:
        print("Token invalid")
        return "Invalid"



# Check leader status function

ind = 0

def get_leader():
    try:
        global ind
        ind = (ind + 1) % 5
        print('trying: ', servers[ind])
        st.session_state['channel'] = grpc.insecure_channel(servers[ind])
        st.session_state['stub'] = LMS_pb2_grpc.LMSStub(st.session_state['channel'])
        request = LMS_pb2.requestPayload()
        response = st.session_state['stub'].get_leader_id(request)
        if response.status == 200:
            st.session_state['leader'] = response.leaderId
        else:
            st.session_state['leader'] = 'none'
    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'

# Login Page
def login_page():
    st.title("Login")
    username = st.text_input("Username", key={'login'})
    password = st.text_input("Password", type="password", key={'login_p'})
    if st.button("Login"):
        try:
            if username == '':
                st.error("Username cannot be empty") 
                return
            if password == '':
                st.error("Password cannot be empty") 
                return

            request = LMS_pb2.loginPayload(username=username, password=password)
            response = st.session_state['stub'].login(request)
            
            if response.status == 200:
                st.session_state['token'] = response.token
                with placeholder.container():
                    st.success("Login successful!")
                    time.sleep(1)
                    handle_state_change("gotToken")
                    return
            else:
                st.error("Login failed. Please try again.")
        except grpc.RpcError as e:
            st.session_state['leader'] = 'none'
            handle_state_change("leaderCrashed")
            return


# Home Page
def home_page():
    st.title("LMS - Distributed Sysyem")

    if st.session_state['role'] == 'student':
        st.subheader(f"Welcome {st.session_state['username']} - STUDENT")
        
        # Navigation bar for Course Materials, Assignments, Queries
        nav_choice = st.sidebar.radio("Navigate", ["Course Materials", "Assignments", "Queries","LLM", "Grades", "Logout"])
        
        if nav_choice == "Course Materials":
            display_course_materials()
        elif nav_choice == "Assignments":
            display_assignments()
        elif nav_choice == "Queries":
            display_queries()
        elif nav_choice == "LLM":
            llm()
        elif nav_choice == "Grades":
            display_grades()
        elif nav_choice == "Logout":
            handle_state_change("logout")

    elif st.session_state['role'] == 'instructor':
        st.subheader(f"Welcome {st.session_state['username']} - INSTRUCTOR")
        
        nav_choice = st.sidebar.radio("Navigate", ["Course Materials", "Assignments", "Queries", "Logout"])

        if nav_choice == "Course Materials":
            inst_display_course_materials()
        elif nav_choice == "Assignments":
            inst_display_assignments()
        elif nav_choice == "Queries":
            inst_display_queries()
        elif nav_choice == "Grades":
            display_grades()
        elif nav_choice == "Logout":
            handle_state_change("logout")

def bytes_to_obj(bytes_data):
    if bytes_data:
        obj_str = bytes_data.decode('utf-8')
        obj = json.loads(obj_str)
        return obj
    return []

def llm():
    # Add custom CSS for consistent styling
    st.markdown(
        """
        <style>
            .query-card {
                background-color: #f9f9f9;
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 16px;
                margin-bottom: 20px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            }
            .query-title {
                font-size: 16px;
                font-weight: bold;
                color: #333;
                margin-bottom: 8px;
            }
            .query-body {
                font-size: 14px;
                color: #555;
                margin-bottom: 10px;
            }
            .query-answer {
                font-size: 14px;
                color: #007bff;
                font-style: italic;
                margin-top: 10px;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    try:
        st.header("LLM Model Logs")
        
        def fetchD():
            request = LMS_pb2.llmLogReq(token=st.session_state['token'])
            response = st.session_state['stub'].get_llm_logs(request)
            logs = bytes_to_obj(response.data)
            querieslist.empty()
            with querieslist.container():
                for log in reversed(logs):
                    st.markdown(
                        f"""
                        <div class="query-card">
                            <div class="query-title"><strong>Query:</strong> {log['question']}</div>
                            <div class="query-answer"><strong>Answer:</strong> {log['answer'] if log['answer'] else "No answer available."}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

        fetchD()

        query_input = st.text_input("Enter your query")
        if st.button("Submit Query", key={"llm"}):
            if query_input.strip():
                request = LMS_pb2.llmRequest(question=query_input.strip(), token=st.session_state['token'])
                response = None
                with st.spinner('Generating answer...'):
                    response = st.session_state['stub'].llm(request)
                if response.answer:
                    st.success("Answer generated successfully!")
                    time.sleep(3)
                    fetchD()
                else:
                    st.error("Failed to submit query.")
            else:
                st.error("Query cannot be empty!")

    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")

    

def inst_display_queries():
    try:
        def fetch_and_display_queries():
            # Request to fetch the queries
            request = LMS_pb2.getQueriesRequest(token=st.session_state['token'])
            response = st.session_state['stub'].get_queries(request)
            queries = response.queries

            # Empty the current container before adding new content
            querieslist.empty()

            # Define custom styling (CSS)
            st.markdown("""
                <style>
                    /* Style for each query card */
                    .query-card {
                        background-color: #f9f9f9;
                        border-radius: 8px;
                        border: 1px solid #000;
                    }

                    /* Style for all paragraph tags with left margin */
                    .query-card p {
                        margin-left: 10px;
                        font-size: 16px;
                        color: #000;
                    }

                    /* Style for query title (Posted By) */
                    .query-title {
                        font-size: 16px;
                        font-weight: bold;
                        color: #000;
                    }

                    /* Style for TextInput labels */
                    .stTextInput label {
                        font-size: 14px;
                        font-weight: bold;
                        color: #000;
                    }

                    /* Style for submit button */
                    .stButton>button {
                        background-color: #000;
                        color: white;
                        border-radius: 3px;
                        font-size: 12px;
                        cursor: pointer;
                        border: none;
                        width: auto;
                        margin-bottom: 2px;
                    }

                    /* Style for submit button on hover */
                    .stButton>button:hover {
                        background-color: #333;
                    }

                    /* Add some space between input fields and buttons */
                    .stTextInput, .stButton {
                    }
                </style>
            """, unsafe_allow_html=True)

            # Display queries
            with querieslist.container():
                for key, query in queries.items():

                    # If the query already has an answer, display it
                    if query.answer != '':
                        st.markdown(f"""
                            <div class="query-card">
                                <p class="query-title">Posted By: {query.postedBy}</p>
                                <p><strong>Query:</strong> {query.query}</p>
                                <p><strong>Answer:</strong> {query.answer}</p>
                            </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown(f"""
                            <div class="query-card">
                                <p class="query-title">Posted By: {query.postedBy}</p>
                                <p><strong>Query:</strong> {query.query}</p>
                        """, unsafe_allow_html=True)

                        query_input = st.text_input("Enter your Answer", key=f"input_{key}")
                        if st.button(f"Submit Answer", key=f"submit_{key}"):
                            if query_input != '':
                                st.write(f"Submitted Answer: {query_input}")
                                request = LMS_pb2.answerQueryPayload(queryId=key, answer=query_input, token=st.session_state['token'])
                                response = st.session_state['stub'].answer_query(request)
                                if response.status == 200:
                                    st.success("Query submitted successfully!")
                                    time.sleep(2)
                                    st.rerun()
                                else:
                                    st.error("Failed to submit answer.")
                            else:
                                st.error("Answer cannot be empty!")
                        st.markdown("</div>", unsafe_allow_html=True)
                    
                    # Add a horizontal line to separate queries
                    st.markdown("<hr style='border: 1px solid #000;'>", unsafe_allow_html=True)

        # Display header and fetch queries
        st.header("Queries")
        fetch_and_display_queries()

    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")


def inst_display_assignments():
    st.header("Assignments")

    # Initialize session state variables for managing assignment views and selected assignment
    if "selected_assignment_id" not in st.session_state:
        st.session_state.selected_assignment_id = None
    if "grading_student" not in st.session_state:
        st.session_state.grading_student = None

    try:
        with st.container():
            def assignmentsList():
                # Custom Styling
                st.markdown("""
                    <style>
                    .stButton>button {
                        background-color: #000;
                        color: white;
                        padding: 14px;
                        width: 23%;
                        height: 5px;
                        border-radius: 5px;
                        font-size: 8px;
                        cursor: pointer;
                        border: none;
                    }
                    .stButton>button:hover {
                        background-color: #333;
                    }
                    .assignment-card {
                        background-color: #fff;
                        padding: 15px;
                        border-radius: 8px;
                        margin-top: 10px;
                        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
                        border: 1px solid #ddd;
                    }
                    .assignment-title {
                        font-size: 18px;
                        font-weight: bold;
                        color: #000;
                    }
                    .assignment-id {
                        color: #777;
                    }
                    .stTextInput label {
                        font-size: 16px;
                        font-weight: bold;
                        color: #000;
                    }
                    </style>
                """, unsafe_allow_html=True)

                # Upload new assignment in col1
                uploaded_file = st.file_uploader("Upload New Assignment", type=["pdf", "docx", "txt"], label_visibility="collapsed")
                assignment_name = st.text_input("Enter Assignment Name")
            

                if st.button("Post Assignment", use_container_width=True):
                    if uploaded_file is not None and assignment_name:
                        file_bytes = uploaded_file.read()
                        submit_request = LMS_pb2.postAssignmentPayload(
                            assignmentName=assignment_name,
                            token=st.session_state['token'],
                            data=file_bytes,
                            filename=uploaded_file.name
                        )
                        submit_response = st.session_state['stub'].post_assignment(submit_request)
                        if submit_response.status == 200:
                            st.success(f"Assignment '{assignment_name}' posted successfully!", icon="✅")
                            st.rerun()
                            
                        else:
                            st.error(f"Failed to submit assignment '{assignment_name}'.", icon="❌")
                    else:
                        st.error("Please upload a file and enter the assignment name.", icon="⚠️")
                
                st.markdown(f"""
                            <hr style="border: 1px solid #ddd; margin: 10px 0;">
                        """, unsafe_allow_html=True)

                # Debug print for development purposes
                print("In assignmentsList function")

                # Request to get assignments list
                request = LMS_pb2.getAssignmentsRequest(token=st.session_state['token'])
                response = st.session_state['stub'].get_assignments(request)

                # Display each assignment with a "View Submissions" button
                for id, name in response.assignments.items():
                    # Create a styled card for each assignment
                    st.markdown(f"""
                        <p class="assignment-title">{name}</p>
                        <p class="assignment-id">ID: {id}</p>
                    """, unsafe_allow_html=True)

                    # Button logic
                    if st.button(f"View Submissions", key=f"view_{id}"):
                        st.session_state.selected_assignment_id = id
                        st.rerun()

                    st.markdown(f"""
                        <hr style="border: 1px solid #ddd; margin: 10px 0;">
                    """, unsafe_allow_html=True)
                        

            # Function to display submissions for a specific assignment
            def displaySubmissions(assignment_id):

                # Button to go back to assignments list
                if st.button("Back to Assignments List"):
                    st.session_state.selected_assignment_id = None
                    st.rerun()

                req = LMS_pb2.getAllSubmissionsRequest(assignmentId=assignment_id, token=st.session_state['token'])
                res = st.session_state['stub'].get_all_submissions_of_one_assignment(req)
                
                st.subheader(f"Submissions for Assignment ID: {assignment_id}")
                
                # Use a container to structure the display better
                for student_username, status in res.submissions.items():
                    # Card-like structure for each student submission
                    with st.expander(f"Student: {student_username} - Status: {status}"):
                        # Display student info and download button
                        st.markdown(f"**Student Username:** {student_username}")
                        st.markdown(f"**Status:** {status}")
                        
                        # Download submission button
                        downloadReq = LMS_pb2.downloadSubmissionRequest(
                            studentUsername=student_username,
                            assignmentId=assignment_id,
                            token=st.session_state['token']
                        )
                        downloadRes = st.session_state['stub'].download_submission(downloadReq)
                        if f"show_pdf_{student_username+assignment_id}" not in st.session_state:
                            st.session_state[f"show_pdf_{student_username+assignment_id}"] = False

                        if st.button("📄 View Submission", key=f"view_btn_{assignment_id+student_username}"):
                            st.session_state[f"show_pdf_{student_username+assignment_id}"] = True
                            st.rerun()

                        file_extension = downloadRes.filename.split('.')[1]

                        if st.session_state[f"show_pdf_{student_username+assignment_id}"]:
                            # Close button
                            if st.button("Close Viewer", key=f"close_btn_{assignment_id+student_username}"):
                                st.session_state[f"show_pdf_{student_username+assignment_id}"] = False
                                st.rerun()

                            if file_extension == "pdf":
                                # Convert file bytes to Base64 for rendering
                                pdf_base64 = base64.b64encode(downloadRes.data).decode('utf-8')
                                pdf_display = f"""
                                <iframe src="data:application/pdf;base64,{pdf_base64}" 
                                        width="100%" height="600px" style="border: none;">
                                </iframe>
                                """
                                st.markdown(pdf_display, unsafe_allow_html=True)

                            elif file_extension == "txt":
                                # Display plain text file
                                text_content = downloadRes.data.decode("utf-8")
                                st.text_area("Text File Content", text_content, height=400, key={assignment_id+student_username+'textview'})

                            elif file_extension == "docx":
                                # Display Word document
                                doc = Document(io.BytesIO(downloadRes.data))
                                st.write("Word Document Content:")
                                for paragraph in doc.paragraphs:
                                    st.write(paragraph.text)

                            else:
                                st.error("Unsupported file format for viewing.")


                        st.download_button(
                            label="Download Submission",
                            data=downloadRes.data,
                            file_name=downloadRes.filename,
                            mime="application/octet-stream",
                            use_container_width=True
                        )
                        
                        # Grade input section
                        grade = st.number_input(
                            label="Enter grade (0-10)",
                            min_value=0,
                            max_value=10,
                            step=1,
                            value=0,
                            key=f"grade_{student_username}",
                            help="Enter a grade between 0 and 10"
                        )
                        
                        # Grade submission button
                        grade_button = st.button(f"Grade {student_username}", key=f"grade_btn_{student_username}")
                        
                        if grade_button:
                            if 0 < grade <= 10:
                                gradeReq = LMS_pb2.gradeAssignmentRequest(
                                    studentUsername=student_username,
                                    assignmentId=assignment_id,
                                    grade=grade,
                                    token=st.session_state['token']
                                )
                                gradeRes = st.session_state['stub'].grade_submission(gradeReq)
                                
                                if gradeRes.status == 200:
                                    st.success(f"Graded {student_username} successfully!")
                                else:
                                    st.error(f"Failed to submit grade for {student_username}. Please try again.")
                            else:
                                st.error("Grade must be between 0 and 10. Please enter a valid grade.")
                        
                        st.markdown("---")

            if st.session_state.selected_assignment_id:
                displaySubmissions(st.session_state.selected_assignment_id)
            else:
                assignmentsList()

            

    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")


def inst_display_course_materials():
    try:
        def fetch_and_display():
            # Fetch all course materials
            request = LMS_pb2.getAllCourseMaterialsRequest(token=st.session_state['token'])
            response = st.session_state['stub'].get_all_course_materials(request)

            # Empty the current container before adding new content
            querieslist.empty()

            # Display each course material
            with querieslist.container():
                for id in response.materials:
                    req = LMS_pb2.downloadCourseMaterialRequest(courseMaterialId=id, token=st.session_state['token'])
                    res = st.session_state['stub'].download_course_material(req)
                    
                    st.markdown("<hr style='border: 1px solid #ddd; margin: 0;'>", unsafe_allow_html=True)
                    # Display course material details
                    st.write(f"#### {response.materials[id]}")
                    if f"show_pdf_{id}" not in st.session_state:
                        st.session_state[f"show_pdf_{id}"] = False

                    if st.button("📄 View Material", key=f"view_btn_{id}"):
                        st.session_state[f"show_pdf_{id}"] = True
                        st.rerun()

                    file_extension = res.filename.split('.')[1]

                    if st.session_state[f"show_pdf_{id}"]:
                        # Close button
                        if st.button("Close Viewer", key=f"close_btn_{id}"):
                            st.session_state[f"show_pdf_{id}"] = False
                            st.rerun()

                        if file_extension == "pdf":
                            # Convert file bytes to Base64 for rendering
                            pdf_base64 = base64.b64encode(res.data).decode('utf-8')
                            pdf_display = f"""
                            <iframe src="data:application/pdf;base64,{pdf_base64}" 
                                    width="100%" height="600px" style="border: none;">
                            </iframe>
                            """
                            st.markdown(pdf_display, unsafe_allow_html=True)

                        elif file_extension == "txt":
                            # Display plain text file
                            text_content = res.data.decode("utf-8")
                            st.text_area("Text File Content", text_content, height=400, key={id+'textview'})

                        elif file_extension == "docx":
                            # Display Word document
                            doc = Document(io.BytesIO(res.data))
                            st.write("Word Document Content:")
                            for paragraph in doc.paragraphs:
                                st.write(paragraph.text)

                        else:
                            st.error("Unsupported file format for viewing.")


                    st.download_button(label="Download", data=res.data, file_name=res.filename, key={id+'down'})


        # Define custom styling (CSS)
        st.markdown("""
            <style>
                /* Style for each material card */
                .query-card {
                    background-color: #f9f9f9;
                    border-radius: 8px;
                    border: 1px solid #000;
                    padding: 15px;
                }

                /* Style for all paragraph tags with left margin */
                .query-card p {
                    margin-left: 10px;
                    font-size: 14px;
                    color: #000;
                }

                /* Style for material title */
                .query-title {
                    font-size: 16px;
                    font-weight: bold;
                    color: #000;
                }

                /* Style for file uploader and text input labels */
                .stTextInput label, .stFileUploader label {
                    font-size: 14px;
                    font-weight: bold;
                    color: #000;
                }

                /* Style for submit button */
                .stButton>button {
                    background-color: #000;
                    color: white;
                    border-radius: 3px;
                    font-size: 12px;
                    cursor: pointer;
                    border: none;
                    width: auto;
                }

                /* Style for submit button on hover */
                .stButton>button:hover {
                    background-color: #333;
                }

                /* Add some space between input fields and buttons */
                
            </style>
        """, unsafe_allow_html=True)

        # Display header and input fields
        st.header("Course Materials")
        uploaded_file = st.file_uploader("Upload Course Material", type=["pdf", "docx", "txt"], key="file_uploader")
        material_name = st.text_input("Enter material name")

        # Handle file upload and material submission
        if st.button("Post Material", key="submit_button"):
            if uploaded_file is None:
                st.error("Please select a file")
            elif material_name == '':
                st.error("Material name cannot be emmpty")
            else:
                file_bytes = uploaded_file.read()
                submit_request = LMS_pb2.postCourseMaterial(
                    materialName=material_name,
                    token=st.session_state['token'],
                    data=file_bytes,
                    filename=uploaded_file.name
                )
                submit_response = st.session_state['stub'].post_course_material(submit_request)
                if submit_response.status == 200:
                    st.success(f"Material - {material_name} posted successfully!")
                    time.sleep(2)
                    material_name = ''
                    uploaded_file = None
                    st.rerun()
                else:
                    st.error("Failed to post material.")

        # Fetch and display course materials
        fetch_and_display()

    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")

        
def display_grades():
    try:
        request = LMS_pb2.retrieveGradesRequest(token=st.session_state['token'])
        response = st.session_state['stub'].retrieve_grades(request)
        if response.status == 200:
            for id in response.grades:
                # Applying card style for better display
                st.markdown(
                    f"""
                    <div class="query-card" style="background-color: #f9f9f9; padding: 15px; margin-bottom: 15px; border-radius: 8px; box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);">
                        <div class="query-title" style="font-size: 16px; font-weight: bold; color: #333;">{id}</div>
                        <div class="query-answer" style="font-size: 14px; color: #555; margin-top: 10px;">
                            <strong>Grade:</strong> {response.grades[id]}
                        </div>
                    </div>
                    <hr style="border: 1px solid #ddd; margin: 10px 0;">
                    """,
                    unsafe_allow_html=True
                )
                
    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")



def display_course_materials():
    st.header("Course Materials")
    try:
        def fetch_and_display():
            # Fetch all course materials
            request = LMS_pb2.getAllCourseMaterialsRequest(token=st.session_state['token'])
            response = st.session_state['stub'].get_all_course_materials(request)

            querieslist.empty()

            with querieslist.container():
                for id in response.materials:
                    req = LMS_pb2.downloadCourseMaterialRequest(courseMaterialId=id, token=st.session_state['token'])
                    res = st.session_state['stub'].download_course_material(req)
                    
                    st.markdown("<hr style='border: 1px solid #ddd; margin: 0;'>", unsafe_allow_html=True)
                    # Display course material details
                    st.write(f"#### {response.materials[id]}")
                    if f"show_pdf1_{id}" not in st.session_state:
                        st.session_state[f"show_pdf1_{id}"] = False

                    if st.button("📄 View Material", key=f"view_btn1_{id}"):
                        st.session_state[f"show_pdf1_{id}"] = True
                        st.rerun()

                    file_extension = res.filename.split('.')[1]

                    if st.session_state[f"show_pdf1_{id}"]:
                        # Close button
                        if st.button("Close Viewer", key=f"close_btn1_{id}"):
                            st.session_state[f"show_pdf1_{id}"] = False
                            st.rerun()

                        if file_extension == "pdf":
                            # Convert file bytes to Base64 for rendering
                            pdf_base64 = base64.b64encode(res.data).decode('utf-8')
                            pdf_display = f"""
                            <iframe src="data:application/pdf;base64,{pdf_base64}" 
                                    width="100%" height="600px" style="border: none;">
                            </iframe>
                            """
                            st.markdown(pdf_display, unsafe_allow_html=True)

                        elif file_extension == "txt":
                            # Display plain text file
                            text_content = res.data.decode("utf-8")
                            st.text_area("Text File Content", text_content, height=400, key={id+'textview1'})

                        elif file_extension == "docx":
                            # Display Word document
                            doc = Document(io.BytesIO(res.data))
                            st.write("Word Document Content:")
                            for paragraph in doc.paragraphs:
                                st.write(paragraph.text)

                        else:
                            st.error("Unsupported file format for viewing.")
                    st.download_button(label="Download", data=res.data, file_name=res.filename, key={id+'1'})
        fetch_and_display()
    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")

def display_assignments():
    st.header("Assignments")

    # Add custom CSS for consistent styling
    st.markdown(
        """
        <style>
            .assignment-card {
                background-color: #f9f9f9;
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 16px;
                margin-bottom: 20px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            }
            .assignment-title {
                font-size: 18px;
                font-weight: bold;
                color: #333;
            }
            .assignment-id {
                font-size: 14px;
                color: #555;
                margin-bottom: 10px;
            }
            .upload-section {
                margin-top: 15px;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    try:
        request = LMS_pb2.getAssignmentsRequest(token=st.session_state['token'])
        response = st.session_state['stub'].get_assignments(request)

        for id in response.assignments:

            st.markdown("<hr style='border: 1px solid #ddd; margin: 0;'>", unsafe_allow_html=True)

            req = LMS_pb2.downloadAssignmentRequest(assignmentId=id, token=st.session_state['token'])
            res = st.session_state['stub'].download_assignment(req)

            # Create a card-like structure for each assignment
            st.markdown(
                f"""
                <div class="">
                    <div class="assignment-title">Assignment: {response.assignments[id]}</div>
                    <div class="assignment-id">ID: {id}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            if f"show_pdf2_{id}" not in st.session_state:
                st.session_state[f"show_pdf2_{id}"] = False

            if st.button("📄 View Assignment", key=f"view_btn2_{id}"):
                st.session_state[f"show_pdf2_{id}"] = True
                st.rerun()

            file_extension = res.filename.split('.')[1]

            if st.session_state[f"show_pdf2_{id}"]:
                # Close button
                if st.button("Close Viewer", key=f"close_btn_{id}"):
                    st.session_state[f"show_pdf2_{id}"] = False
                    st.rerun()

                if file_extension == "pdf":
                    # Convert file bytes to Base64 for rendering
                    pdf_base64 = base64.b64encode(res.data).decode('utf-8')
                    pdf_display = f"""
                    <iframe src="data:application/pdf;base64,{pdf_base64}" 
                            width="100%" height="600px" style="border: none;">
                    </iframe>
                    """
                    st.markdown(pdf_display, unsafe_allow_html=True)

                elif file_extension == "txt":
                    # Display plain text file
                    text_content = res.data.decode("utf-8")
                    st.text_area("Text File Content", text_content, height=400, key={id+'textview2'})

                elif file_extension == "docx":
                    # Display Word document
                    doc = Document(io.BytesIO(res.data))
                    st.write("Word Document Content:")
                    for paragraph in doc.paragraphs:
                        st.write(paragraph.text)

                else:
                    st.error("Unsupported file format for viewing.")

            # Add a download button
            st.download_button(
                label="Download",
                data=res.data,
                file_name=res.filename,
                key=f"download_button_{id}",
            )

            # Add file uploader and submission button
            uploaded_file = st.file_uploader(
                f"Upload Submission for ID: {id}",
                type=["pdf", "docx", "txt"],
                key=f"file_uploader_{id}",
            )

            if uploaded_file is not None:
                if st.button(f"Submit Assignment for ID: {id}", key=f"submit_button_{id}"):
                    file_bytes = uploaded_file.read()
                    submit_request = LMS_pb2.submitAssignment(
                        assignmentId=id,
                        token=st.session_state['token'],
                        data=file_bytes,
                        filename=uploaded_file.name,
                    )
                    submit_response = st.session_state['stub'].submit_assignment(submit_request)
                    if submit_response.status == 200:
                        st.success(f"Assignment {id} submitted successfully!")
                    else:
                        st.error(f"Failed to submit assignment {id}.")

    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")

def display_queries():
    # Add custom CSS for consistent styling
    st.markdown(
        """
        <style>
            .query-card {
                background-color: #f9f9f9;
                border: 1px solid #ddd;
                border-radius: 8px;
                padding: 16px;
                margin-bottom: 20px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            }
            .query-title {
                font-size:20px;
                font-weight: bold;
                color: #333;
                margin-bottom: 8px;
            }
            .query-body {
                font-size: 18px;
                color: #555;
                margin-bottom: 10px;
            }
            .query-answer {
                font-size: 18px;
                color: #007bff;
                font-style: italic;
                margin-top: 10px;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    try:
        def fetch_and_display_queries():
            request = LMS_pb2.getQueriesRequest(token=st.session_state['token'])  # Adjust according to your proto file
            response = st.session_state['stub'].get_queries(request)
            queries = response.queries
            querieslist.empty()
            with querieslist.container():
                for key in queries:
                    query = queries[key]
                    st.markdown(
                        f"""
                        <div class="query-card">
                            <div class="query-title">Posted by: {query.postedBy}</div>
                            <div class="query-body"><strong>Query:</strong> {query.query}</div>
                            <div class="query-answer"><strong>Answer:</strong> {query.answer if query.answer else "No answer yet."}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

        st.header("Queries")
        query_input = st.text_input("Enter your query", key={'ques'})

        if st.button("Submit Query"):
            if query_input.strip():
                request = LMS_pb2.postQueryPayload(query=query_input.strip(), token=st.session_state['token'])
                response = st.session_state['stub'].post_query(request)
                if response.status == 200:
                    st.success(f"Query submitted successfully!")
                    time.sleep(2)
                    fetch_and_display_queries()
                else:
                    st.error("Failed to submit query.")
            else:
                st.error("Query cannot be empty!")

        fetch_and_display_queries()

    except grpc.RpcError as e:
        st.session_state['leader'] = 'none'
        handle_state_change("leaderCrashed")

def handle_state_change(change):
        placeholder.empty()
        with placeholder.container():
            if change == "gotLeader":
                display_leader_status()  
                print(st.session_state['leader'])
                st.session_state['channel']  = grpc.insecure_channel(st.session_state['leader'])
                st.session_state['stub'] = LMS_pb2_grpc.LMSStub(st.session_state['channel'])
                if st.session_state['token']:
                    home_page()
                else:
                    login_page()

            elif change == "leaderCrashed":
                display_leader_status()  
                st.warning("Waiting for Leader...")
                while(st.session_state['leader'] == 'none'):
                    get_leader()
                    time.sleep(0.5)
                handle_state_change("gotLeader")

            elif change == "gotToken":
                verify_token(st.session_state['token'])
                home_page()
            
            elif change == "logout":
                st.session_state['token'] = None
                login_page()

# Main logic for the app
with placeholder.container():
    if st.session_state['leader'] == 'none':
        st.warning("Waiting for Leader...")
        while(st.session_state['leader'] == 'none'):
            get_leader()
            time.sleep(0.5)
        handle_state_change("gotLeader")
    else:
        if st.session_state['token'] is None:
            login_page()
        else:
            home_page()

